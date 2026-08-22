"""Unit tests for src/extract.py -- real temp PDF/DOCX files, no network."""

import pytest
from docx import Document
from fpdf import FPDF

from src.errors import FileNotFoundPipelineError, NoExtractableTextError, UnsupportedFileTypeError
from src.extract import extract_text


def test_missing_file_raises():
    with pytest.raises(FileNotFoundPipelineError):
        extract_text("this_file_absolutely_does_not_exist_12345.pdf")


def test_unsupported_extension_raises(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("hello", encoding="utf-8")
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(str(p))


def test_directory_path_raises_file_not_found(tmp_path):
    # A directory is not a file -- must not be treated as "found".
    with pytest.raises(FileNotFoundPipelineError):
        extract_text(str(tmp_path))


def test_empty_docx_raises_no_extractable_text(tmp_path):
    p = tmp_path / "empty.docx"
    Document().save(str(p))
    with pytest.raises(NoExtractableTextError):
        extract_text(str(p))


def test_docx_with_only_whitespace_raises_no_extractable_text(tmp_path):
    p = tmp_path / "whitespace.docx"
    doc = Document()
    doc.add_paragraph("   ")
    doc.add_paragraph("\t\t")
    doc.save(str(p))
    with pytest.raises(NoExtractableTextError):
        extract_text(str(p))


def test_docx_with_real_text_extracts_it(tmp_path):
    p = tmp_path / "real.docx"
    doc = Document()
    doc.add_paragraph("First paragraph of real content.")
    doc.add_paragraph("Second paragraph follows here.")
    doc.save(str(p))
    text = extract_text(str(p))
    assert "First paragraph" in text
    assert "Second paragraph" in text


def test_docx_paragraphs_separated_by_blank_line(tmp_path):
    # Paragraph-chunking strategy depends on this.
    p = tmp_path / "structured.docx"
    doc = Document()
    doc.add_paragraph("Alpha.")
    doc.add_paragraph("Beta.")
    doc.save(str(p))
    text = extract_text(str(p))
    assert "\n\n" in text


def test_docx_table_cells_are_extracted(tmp_path):
    p = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CellOne"
    table.rows[0].cells[1].text = "CellTwo"
    doc.save(str(p))
    text = extract_text(str(p))
    assert "CellOne" in text
    assert "CellTwo" in text


def test_pdf_with_real_text_extracts_it(tmp_path):
    p = tmp_path / "real.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Extractable PDF content here")
    pdf.output(str(p))
    text = extract_text(str(p))
    assert "Extractable PDF content" in text


def test_pdf_multi_page_concatenates_pages(tmp_path):
    p = tmp_path / "multipage.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Page one content marker")
    pdf.add_page()
    pdf.cell(0, 10, "Page two content marker")
    pdf.output(str(p))
    text = extract_text(str(p))
    assert "Page one content marker" in text
    assert "Page two content marker" in text


def test_extraction_collapses_excess_whitespace(tmp_path):
    p = tmp_path / "messy.docx"
    doc = Document()
    doc.add_paragraph("Text   with     excess     spaces.")
    doc.save(str(p))
    text = extract_text(str(p))
    assert "     " not in text  # 5+ consecutive spaces should be collapsed


def test_case_insensitive_extension(tmp_path):
    p = tmp_path / "upper.DOCX"
    doc = Document()
    doc.add_paragraph("Content in an uppercase-extension file.")
    doc.save(str(p))
    text = extract_text(str(p))
    assert "Content in an uppercase-extension file." in text
