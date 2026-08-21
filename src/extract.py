"""Text extraction for PDF and DOCX files."""

from __future__ import annotations

import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from .errors import (
    FileNotFoundPipelineError,
    NoExtractableTextError,
    UnsupportedFileTypeError,
)

SUPPORTED_SUFFIXES = {".pdf", ".docx"}


def extract_text(path: str | Path) -> str:
    """Extract and clean text from a PDF or DOCX file.

    Raises:
        FileNotFoundPipelineError: the path does not exist.
        UnsupportedFileTypeError: the extension is not .pdf or .docx.
        NoExtractableTextError: the file parses but yields no usable text
            (e.g. a scanned/image-only PDF with no text layer).
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundPipelineError(f"File not found: {p}")

    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    if suffix == ".pdf":
        raw = _extract_pdf(p)
    else:
        raw = _extract_docx(p)

    cleaned = _clean_text(raw)
    if not cleaned:
        raise NoExtractableTextError(
            f"No extractable text found in {p.name}. "
            "The file may be a scanned/image-only document with no text layer."
        )
    return cleaned


def _extract_pdf(p: Path) -> str:
    try:
        reader = PdfReader(str(p))
    except Exception as e:  # corrupt/encrypted/malformed PDF
        raise NoExtractableTextError(f"Could not open PDF {p.name}: {e}") from e

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise NoExtractableTextError(
                f"{p.name} is password-protected and could not be opened."
            )

    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(p: Path) -> str:
    try:
        doc = Document(str(p))
    except Exception as e:
        raise NoExtractableTextError(f"Could not open DOCX {p.name}: {e}") from e

    # Blank line between paragraphs preserves paragraph boundaries for the
    # paragraph-based chunking strategy; DOCX has no other structural marker
    # for "this is a new paragraph" once text is flattened.
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n\n".join(parts)


def _clean_text(text: str) -> str:
    """Collapse whitespace noise from extraction without destroying paragraph breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # collapse runs of spaces/tabs, but keep newlines meaningful for later chunking
    text = re.sub(r"[ \t]+", " ", text)
    # collapse 3+ blank lines down to a single paragraph break
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()
